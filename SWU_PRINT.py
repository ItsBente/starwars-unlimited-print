from docx import Document
from docx.shared import Cm, Inches, Pt
import os
import json
import requests
from PIL import Image
import tkinter as tk
from tkinter import filedialog

# Function to download an image from URL to a local path
def download_image(url, save_path):
    try:
        response = requests.get(url, stream=True)
        if response.status_code == 200:
            with open(save_path, 'wb') as file:
                for chunk in response.iter_content(1024):
                    file.write(chunk)
            return save_path
        else:
            print(f"Failed to download image from {url}. Status code: {response.status_code}")
            return None
    except requests.exceptions.RequestException as e:
        print(f"Error downloading image from {url}: {str(e)}")
        return None

# Function to convert WebP image to PNG
def convert_webp_to_png(image_path):
    try:
        img = Image.open(image_path)
        if img.format == 'webp' or 'png':  # Fixed the format check
            png_image_path = image_path.rsplit('.', 1)[0] + '.png'
            img.save(png_image_path, 'PNG')
            os.remove(image_path)  # Remove the original WebP file
            print("Image Converted:", image_path)
            return png_image_path
        else:
            print("Image not Converted:", image_path)
            return image_path
    except Exception as e:
        print(f"Error converting WebP to PNG: {str(e)}")
        return None

# Function to fetch images based on card IDs from a JSON file
def fetch_images_from_website(json_file_path, image_folder):
    try:
        with open(json_file_path, 'r') as file:
            data = json.load(file)
        
        base_url = "https://swudb.com/cards/SOR/"

        # Ensure the image folder exists
        if not os.path.exists(image_folder):
            os.makedirs(image_folder)
        
        card_ids = set()

        # Collect all unique card IDs
        if data.get('leader'):
            card_ids.add(data['leader']['id'].split('_')[1])
        if data.get('base'):
            card_ids.add(data['base']['id'].split('_')[1])
        if data.get('deck'):
            for card in data['deck']:
                card_ids.add(card['id'].split('_')[1])
        if data.get('sideboard'):
            for card in data['sideboard']:
                card_ids.add(card['id'].split('_')[1])
        
        print(f"Card IDs to process: {card_ids}")

        for card_id in card_ids:
            # Check if PNG file already exists in the folder
            png_filename = os.path.join(image_folder, f"{card_id}.png")
            if not os.path.exists(png_filename):
                # Download regular image if it doesn't exist
                regular_image_url = f"{base_url}{card_id}.png"
                if requests.head(regular_image_url).status_code == 200:
                    print(f"Downloading regular image for card ID {card_id} from URL: {regular_image_url}")
                    img_filename = os.path.join(image_folder, f"{card_id}.webp")
                    download_image(regular_image_url, img_filename)
                    convert_webp_to_png(img_filename)

            # Check if portrait PNG file already exists in the folder
            portrait_png_filename = os.path.join(image_folder, f"{card_id}-portrait.png")
            if not os.path.exists(portrait_png_filename):
                # Download portrait image if it doesn't exist
                portrait_image_url = f"{base_url}{card_id}-portrait.png"
                if requests.head(portrait_image_url).status_code == 200:
                    print(f"Downloading portrait image for card ID {card_id} from URL: {portrait_image_url}")
                    portrait_img_filename = os.path.join(image_folder, f"{card_id}-portrait.webp")
                    download_image(portrait_image_url, portrait_img_filename)
                    convert_webp_to_png(portrait_img_filename)

    except Exception as e:
        print(f"Error fetching images: {str(e)}")

# Function to rotate image if width > height
def rotate_image_if_needed(image_path):
    try:
        img = Image.open(image_path)
        width, height = img.size
        
        if width > height:
            img = img.rotate(90, expand=True)  # Rotate and expand the image to avoid cropping
            rotated_image_path = image_path.rsplit('.', 1)[0] + '_rotated.png'
            img.save(rotated_image_path, 'PNG')
            return rotated_image_path
        else:
            return image_path
    except Exception as e:
        print(f"Error rotating image {image_path}: {str(e)}")
        return None

# Modify add_images_to_word function to integrate image rotation and accept JSON path
def add_images_to_word(image_folder, output_file, json_file_path):
    try:
        doc = Document()
        
        count = 0

        # Set Word document margins (0.5 cm on all sides)
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(0.5)
            section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(0.5)
            section.right_margin = Cm(0.5)

        # Open and load JSON data
        with open(json_file_path, 'r') as file:
            data = json.load(file)

        images = []
        base_url = "https://swudb.com/cards/SOR/"

        # Function to add images to the images list based on count
        def add_images(card_list):
            for card in card_list:
                card_id = card['id'].split('_')[1]
                count = card['count']
                for _ in range(count):
                    images.append(f"{card_id}.png")  # Assuming the file names are {card_id}.png

        # Add leader and base images
        if data.get('leader'):
            images.append(f"{data['leader']['id'].split('_')[1]}.png")
            images.append(f"{data['leader']['id'].split('_')[1]}-portrait.png")
        if data.get('base'):
            images.append(f"{data['base']['id'].split('_')[1]}.png")

        # Add deck and sideboard images
        if data.get('deck'):
            add_images(data['deck'])
        if data.get('sideboard'):
            add_images(data['sideboard'])

        # Calculate number of pages needed (each page contains 3 rows of images)
        num_images_per_page = 9
        num_pages = (len(images) + num_images_per_page - 1) // num_images_per_page

        # Constants for table dimensions
        table_width_cm = 6.75  # Width of the table in cm
        table_height_cm = 8.7  # Height of the table in cm

        # Convert table dimensions from cm to inches
        table_width = Cm(table_width_cm).inches
        table_height = Cm(table_height_cm).inches

        img_width = Cm(table_width_cm - .5).inches
        img_height = Cm(table_height_cm).inches

        # Loop through images and add to document
        for page in range(num_pages):
            # Add a page break for all pages except the first one
            if page > 0:
                doc.add_page_break()

            # Add 3 rows of images per page
            for row in range(3):
                # Add a new table for each row of images
                table = doc.add_table(rows=1, cols=3)

                # Set table width and height
                table.autofit = False
                table.allow_autofit = False
                table.width = Inches(table_width)
                table.height = Inches(table_height)

                # Calculate available width for each cell
                cell_width = table_width

                # Set column widths
                for cell in table.columns:
                    for col in cell.cells:
                        col.width = Inches(cell_width)

                # Calculate available height for each cell
                cell_height = table_height

                # Set row height
                table.rows[0].height = Inches(cell_height)

                # Add up to 3 images per row
                for col in range(3):
                    index = page * num_images_per_page + row * 3 + col
                    if index < len(images):
                        image_name = images[index]
                        image_path = os.path.join(image_folder, image_name)
                        count = count + 1

                        # Rotate image if needed
                        rotated_image_path = rotate_image_if_needed(image_path)
                        if rotated_image_path:
                            image_path = rotated_image_path
                        else:
                            print(f"Failed to rotate image {image_path}. Skipping.")

                        # Add image to table cell
                        cell = table.cell(0, col)
                        paragraph = cell.paragraphs[0]
                        run = paragraph.add_run()
                        run.add_picture(image_path, width=Inches(img_width), height=Inches(img_height))

                        # Set space after each paragraph to 0 pt
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.space_after = Pt(0)

                        # Check and add portrait image if exists
                        portrait_image_path = image_path.replace('.png', '-portrait.png')
                        if os.path.exists(portrait_image_path):
                            if col == 2:
                                table = doc.add_table(rows=1, cols=3)
                                table.autofit = False
                                table.allow_autofit = False
                                table.width = Inches(table_width)
                                table.height = Inches(table_height)
                                cell_width = table_width
                                for cell in table.columns:
                                    for col in cell.cells:
                                        col.width = Inches(cell_width)
                                cell_height = table_height
                                table.rows[0].height = Inches(cell_height)
                                col = 0
                            else:
                                col += 1

                            cell = table.cell(0, col)
                            paragraph = cell.paragraphs[0]
                            run = paragraph.add_run()
                            run.add_picture(portrait_image_path, width=Inches(img_width), height=Inches(img_height))
                            paragraph_format = paragraph.paragraph_format
                            paragraph_format.space_after = Pt(0)

        doc.save(output_file)
        print(count)
        print(f'Document saved as {output_file}')
        

    except Exception as e:
        print(f"Error adding images to Word document: {str(e)}")

# Function to create a popup window for selecting JSON file
def select_json_file():
    root = tk.Tk()
    root.withdraw()  # Hide the main window
    file_path = filedialog.askopenfilename(filetypes=[("JSON files", "*.json")])
    return file_path

# Function to create a popup window for selecting image folder
def select_image_folder():
    root = tk.Tk()
    root.withdraw()  # Hide the main window
    folder_path = filedialog.askdirectory()
    return folder_path

# Function to create a popup window for selecting output file
def select_output_file():
    root = tk.Tk()
    root.withdraw()  # Hide the main window
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    return file_path

# Example usage
if __name__ == "__main__":
    # Prompt user to select JSON file using popup window
    print("Select json-File to Import.")
    json_file_path = select_json_file()
    if json_file_path:
        # Prompt user to select image folder using popup window
        print("Select the Folder where Images are allready downloaded or you want to download the files to.")
        image_folder = select_image_folder()
        if image_folder:
            # Prompt user to select output file using popup window
            print("Select the name for the output-file.")
            output_file = select_output_file()
            if output_file:
                print("Fetching card IDs and downloading images...")
                fetch_images_from_website(json_file_path, image_folder)
                print("Adding images to Word document...")
                add_images_to_word(image_folder, output_file, json_file_path)
            else:
                print("No output file selected.")
        else:
            print("No image folder selected.")
    else:
        print("No JSON file selected.")
